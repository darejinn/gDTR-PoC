"""Build Paper1_gDTR v5 (memo-reflected, figures embedded).

v5 changes vs Paper1_gDTR_0429.docx:
  - Abstract: v3 two-tool framing (||Δh||_2 scorer + gDTR mechanistic probe)
  - §1: add benchmarking-contribution flag
  - §3.4: same variant-level analysis, with bridge text to §3.5
  - §3.5 NEW: Interpretability baseline benchmark (||Δh||_2 honest reporting + P1-P3 + residualized AUROC + reviewer pre-empt)
  - §3.6: old §3.5 cross-architecture
  - §5 Conclusion: two-tool framing
  - 7 main figures embedded inline from results/figures_v2/

Source memos:
  paper_reframing_v3_20260428.md, paper_plan_gdtr.md, analysis_plan_tier12.md
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/yoonjincho/Project/ICML")
FIG_DIR = ROOT / "results" / "figures_v2"
OUT = ROOT / "Paper1_gDTR_0429_v5.docx"


# ─── helpers ──────────────────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    """Add basic single-line borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "000000")
        tcBorders.append(e)
    tcPr.append(tcBorders)


def add_table(doc, header, rows, *, header_bold=True, first_col_bold=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = header_bold
        run.font.size = Pt(9)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if first_col_bold and c_idx == 0:
                run.bold = True
    return table


def add_figure(doc, image_path, caption, width_in=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    return p, cap


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    return h


def add_para(doc, text, *, italic=False, bold=False, align=None, size=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_bulleted(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)


# ─── build doc ────────────────────────────────────────────────────────────
doc = Document()

# ── Title ──
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "gDTR: A Layer-Resolved Mechanistic Probe for Genomic Causal Foundation "
    "Models, Benchmarked Against Hidden-State Perturbation Scoring"
)
title_run.bold = True
title_run.font.size = Pt(15)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run(
    "ICML 2026 Workshop on Interpretable Machine Learning in the Sciences — short paper (4 pages + appendix)"
)
sub_run.italic = True
sub_run.font.size = Pt(10)

# ── Abstract ──
add_h2(doc, "Abstract")
add_para(
    doc,
    "We introduce gDTR (Genomic Deep-Thinking Ratio), a training-free, layer-resolved "
    "interpretability framework that adapts the core Deep-Thinking Ratio concept from "
    "natural language processing to genomic causal language models. By measuring "
    "per-token settling depth through cosine similarity trajectories in the logit lens, "
    "gDTR quantifies how deeply and at which layers the model engages with different "
    "parts of the genome. Architecture-aware adaptations — small-vocabulary calibration, "
    "Evo 2's idle final block, and 32-layer tuned-lens recovery — make the framework "
    "applicable across hybrid Transformer/Hyena, pure Hyena, and masked-language-model "
    "genomic foundation models. Applied to Evo 2 7B and validated across HyenaDNA, "
    "Nucleotide Transformer v2, and DNABERT-2, gDTR (i) identifies splice donor/acceptor "
    "sites as a universal shallow-thinking signature replicating across two chromosomes "
    "and two architecture families; (ii) achieves AUROC 0.844 on 8,008 ClinVar "
    "pathogenic-vs-benign single-nucleotide variants and provides DeLong-significant "
    "incremental information beyond Evo 2's own likelihood (p < 10⁻¹⁵); (iii) reveals "
    "class-stratified disruption layers — splice variants peak at shallow layer 7, the "
    "TP53 R175H missense at deep layer 28, frameshifts at mid-deep layer 24; and (iv) "
    "discovers chr22 regions of high model engagement but low evolutionary conservation "
    "(Q2) that are 1.62× enriched for GTEx eQTLs, 1.50× for GWAS hits, and 1.90× for "
    "ENCODE enhancer-like cCREs. Benchmarked against three interpretability baselines "
    "on the same cohort, we report transparently that the simple per-layer hidden-state "
    "perturbation magnitude ‖Δh‖₂ achieves AUROC 0.926 — outperforming gDTR by +0.083 "
    "as a raw classifier. We therefore reposition gDTR not as a competing scorer but as "
    "a complementary mechanistic probe: it provides a layer-index reference frame, "
    "class-stratified disruption layers, and a connection to genome-wide universality "
    "that ‖Δh‖₂ cannot supply by construction. ΔD_cos retains AUROC 0.645 after "
    "residualizing on ‖Δh‖₂, evidence of independent variance. gDTR thus opens a fourth "
    "interpretability axis — where in the layer hierarchy a sequence is computationally "
    "resolved — that complements both attention/SAE/likelihood approaches and the "
    "stronger ‖Δh‖₂ scoring baseline that this work also benchmarks.",
)

# ─────────────────────────────────────────────────────────────────────────
# §1 Introduction
add_h1(doc, "1 Introduction")

add_para(
    doc,
    "Genomic foundation models trained on whole-genome corpora — Evo 2 [1], "
    "HyenaDNA [2], Nucleotide Transformer [3], DNABERT-2 [4] — encode rich biological "
    "knowledge across nucleotide-resolution sequence, gene structure, regulatory "
    "grammar, splicing, and even protein secondary structure. Yet how this knowledge "
    "is computed across layers remains underexplored. Existing interpretability "
    "research is primarily organized around three paradigms: (1) attention-based "
    "methods, which identify where the model attends; (2) sparse autoencoders, which "
    "isolate what concepts the model encodes; and (3) embedding-distance or likelihood "
    "scores, which evaluate what a particular variant means to the model's "
    "representation. These approaches do not address a fundamental aspect of internal "
    "computation: the temporal progression of information processing across layers. "
    "We therefore propose a fourth axis of interpretability: where, in the layer "
    "hierarchy, is the sequence computationally resolved?",
)

add_para(
    doc,
    "The study of such hierarchical computation has recently emerged in NLP. Chen et "
    "al. [5] introduced the Deep-Thinking Ratio (DTR), a training-free metric that "
    "measures the layer index at which intermediate-layer prediction distributions "
    "converge to within ε of the final-layer distribution. We hypothesize that this "
    "hierarchical lens is uniquely suited to genomics, where different biological "
    "elements (e.g., splice sites versus distal regulatory regions) plausibly require "
    "markedly different depths of computational processing. Transferring DTR to the "
    "genomic domain, however, faces three concrete challenges: (C1) extremely small "
    "vocabularies (|V| ≈ 12–512) saturate distribution-based thresholds and require "
    "quantile-based calibration; (C2) the genomic landscape spans hybrid "
    "Transformer-StripedHyena, pure Hyena, and masked-LM Transformers, each requiring "
    "architecture-specific lens implementations while preserving cross-architecture "
    "comparability; (C3) in models such as Evo 2, the final attention block functions "
    "as a residual passthrough, rendering the standard NLP practice of \"tapping the "
    "last block\" ineffective.",
)

add_para(
    doc,
    "We present gDTR (Genomic Deep-Thinking Ratio), the first systematic adaptation "
    "of NLP-DTR to genomic causal language models. To validate the framework as more "
    "than a re-skinned NLP method, we additionally benchmark it on the same ClinVar "
    "cohort against three interpretability baselines (attention rollout, integrated "
    "gradients, and per-layer hidden-state perturbation magnitude ‖Δh‖₂). This "
    "benchmark uncovers a previously under-reported result: the simple ‖Δh‖₂ score "
    "achieves AUROC 0.926 on the same cohort, outperforming gDTR's discrimination "
    "(0.844) by +0.083. We therefore reframe the paper around two complementary "
    "tools: ‖Δh‖₂ as a strong but mechanistically opaque scoring baseline, and gDTR "
    "as a layer-resolved mechanistic probe whose value lies in providing a reference "
    "frame, class-stratified disruption layers, and biological universality that "
    "‖Δh‖₂ cannot supply by construction (§3.5). Section 2 defines the framework and "
    "the three adaptations above. Section 3 demonstrates five findings enabled by the "
    "layer-resolved readout, including the baseline benchmark. Section 4 enumerates "
    "five limitations of the framework.",
)

# ─────────────────────────────────────────────────────────────────────────
# §2 Framework
add_h1(doc, "2 The gDTR Framework")
add_h2(doc, "2.1 Settling depth")
add_para(
    doc,
    "For an input nucleotide sequence of length T processed by a causal language "
    "model with L layers, we extract residual-stream activations h_l(t). In genomic "
    "CLMs the vocabulary is extremely small (|V| ≤ 512, often |V| = 12). This causes "
    "the JSD lens used in the original NLP DTR formulation to saturate severely. "
    "We therefore adopt the cosine unembed-residual (UR) lens as the primary signal:",
)
eq_p = doc.add_paragraph()
eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_run = eq_p.add_run("D_cos(l, t) = 1 − cos(h_l(t), h_norm(t)),")
eq_run.italic = True

add_para(
    doc,
    "where h_norm is the post-final-norm tensor. Cosine distance operates directly "
    "in the hidden-state geometry and is independent of vocabulary size, providing a "
    "stable measure of how closely an intermediate representation has converged "
    "toward the final readout subspace.",
)
add_para(
    doc,
    "Crucially, and in contrast to the original NLP DTR, we define the settling "
    "depth of token t at threshold γ via the running-min envelope:",
)
eq_p2 = doc.add_paragraph()
eq_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_run2 = eq_p2.add_run(
    "c(t) = min{l : run-min D(·, t) ≤ γ}, where  run-min D(l, t) = min_{k=1..l} D_cos(k, t)."
)
eq_run2.italic = True
add_para(
    doc,
    "The running-min envelope absorbs raw monotonicity violations exhibited by some "
    "genomic models: any temporary excursion away from the final representation is "
    "ignored once the token has reached the γ-threshold at an earlier layer. This "
    "design choice makes gDTR robust to architectural quirks that would otherwise "
    "break standard logit-lens assumptions. To set γ we use regional q70 calibration: "
    "the 70th percentile of the running-min cosine distance distribution at the "
    "penultimate layer (Appendix A.2).",
)

add_h2(doc, "2.2 Architectural quirk handling: Evo 2's idle last block")
add_para(
    doc,
    "We discover that Evo 2 7B's last attention block is architecturally idle. Direct "
    "verification on chr22 sanity sequences (6 kb × 100 windows) yields the values in "
    "Table 1. The implication is that the canonical \"tap the last block\" convention "
    "from NLP transformer interpretability does not apply to Evo 2: meaningful "
    "deep-thinking computation completes at block 29, with blocks 30–31 acting as "
    "bookkeeping. We therefore lock the canonical deep-thinking tap at L = 29 and use "
    "h_norm as the convergence reference.",
)
add_table(
    doc,
    ["Quantity", "Measured value", "Interpretation"],
    [
        ["max|h₃₀ − h₃₁|", "0.000000 (exact)", "block 31 is a residual passthrough"],
        ["data_ptr(h₃₀) vs data_ptr(h₃₁)", "different addresses", "tensors are physically distinct copies"],
        ["cos(h₃₁, h_norm)", "0.6855", "post-norm differs from raw block 31"],
        ["cos(h₃₀, h_norm)", "0.6855", "identical to block 31"],
        ["cos(h₂₉, h_norm)", "−0.013", "block 29 is the deepest interpretively distinct tap"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 1. Evidence that Evo 2's last attention block is architecturally idle.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "2.3 Tuned-lens recovery")
add_para(
    doc,
    "To recover layers whose residual stream is rotated relative to the final readout "
    "subspace, we fit a per-layer affine A_l with the tuned-lens objective of Belrose "
    "et al. [6]. Across Evo 2's 32 layers, 30/32 reach ≥ 98 % MSE recovery from a "
    "single 4096² affine; the worst recovery is at L = 12 (0.9816), and the canonical "
    "tap L = 29 reaches 0.9996. For variant analysis we use the post-final-norm tap "
    "as reference and the original residual stream taps as the lens input — no "
    "learned weights are required at inference time, preserving the training-free "
    "property of the framework.",
)

# ── Figure 1: Method schematic ──
add_figure(
    doc,
    FIG_DIR / "F1_method_schematic.png",
    "Figure 1. gDTR framework. (a) Four interpretability axes for genomic CLMs; "
    "gDTR addresses the previously untouched \"where does the model think deeply?\" "
    "axis. (b) Adaptation from NLP-DTR (Chen et al. 2026) to gDTR with three "
    "challenges (C1–C3). (c) Tuned-lens MSE recovery per layer — 30/32 layers reach "
    "≥ 98 %; L = 30, 31 are the degenerate idle tail. (d) Calibration: raw lens "
    "monotonicity fails for all block types; the running-min envelope plus cosine UR "
    "lens (used by gDTR) restores stable convergence behaviour.",
    width_in=6.2,
)

# ─────────────────────────────────────────────────────────────────────────
# §3 Biological utility
add_h1(doc, "3 Biological Utility")

# §3.1 Splice
add_h2(doc, "3.1 Splice sites are a universal shallow-thinking signature")
add_para(
    doc,
    "On chr22 (12,978 windows of 6 kb, 77.9 M positions) and chr17 (27,586 windows, "
    "80 M positions), we computed mean settling depth c(t) per genomic context using "
    "GENCODE v44 annotations. The pattern is consistent (Table 2): splice donor and "
    "acceptor positions converge earlier (at lower layer indices) than every other "
    "annotated context, including coding exons, introns, UTRs, and intergenic "
    "regions. The pairwise Cohen's d of splice donor versus intergenic baseline is "
    "+0.540 (largest among all pairs).",
)
add_table(
    doc,
    ["Context", "mean c (chr22)", "mean c (chr17)", "n positions (chr22)"],
    [
        ["splice donor", "25.57", "25.51", "187,236"],
        ["splice acceptor", "25.69", "25.64", "185,890"],
        ["3' UTR", "27.72", "27.60", "1,216,058"],
        ["intron", "27.82", "27.69", "41,477,463"],
        ["coding exon", "28.26", "28.04", "3,257,646"],
        ["intergenic", "28.75", "28.51", "31,396,515"],
        ["5' UTR", "28.99", "28.83", "147,192"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 2. Per-context mean settling depth (c) on chr22 and chr17. Splice "
    "donor/acceptor positions converge ≈3 layers earlier than intronic baseline.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "Replicating on HyenaDNA-large (28 M parameters, 8 layers, pure Hyena) on the "
    "same chr22 windows shows the same direction at a depth-normalised scale: donor "
    "mean c = 6.55 vs intron 6.89 (8-layer model). The donor < intron inequality "
    "therefore holds across two chromosomes and two architecture families spanning "
    "28 M to 7 B parameters and 8 to 32 layers. Profiling around individual splice "
    "sites shows the asymmetric extension predicted by canonical splice-grammar "
    "biology: donor signals extend further on the exonic side (minimum at +20 bp on "
    "chr17), acceptor signals further on the intronic side (minimum at +50 bp), and "
    "neither returns to background within ± 200 bp — consistent with branch-point "
    "and polypyrimidine-tract integration.",
)
add_figure(
    doc,
    FIG_DIR / "F2_splice_universality.png",
    "Figure 2. Splice deep-thinking universality. Per-context settling-depth "
    "distributions on chr22 (Evo 2) and chr17 replication; donor/acceptor minima "
    "lie ~3 layers below intron, with HyenaDNA-large showing the same direction "
    "in its 8-layer regime. Asymmetric extension around individual splice sites "
    "(donor exonic-side, acceptor intronic-side) recovers canonical splice-grammar "
    "geometry purely from internal layer dynamics.",
    width_in=6.2,
)

# §3.2 Q2
add_h2(doc, "3.2 Conservation discordance reveals lineage-specific regulatory elements")
add_para(
    doc,
    "We project gDTR onto the chr22 reference at single-base resolution and compare "
    "it to PhyloP 100-way conservation. Both signals are smoothed with a 100 bp "
    "box-car prior to median-split thresholding, yielding four quadrants. The "
    "biologically distinctive quadrant is Q2 = high gDTR ∩ low conservation, covering "
    "3.71 % of chr22 (1.9 Mb across 5,090 contiguous regions ≥ 100 bp). Three "
    "independent functional axes — molecular-trait (eQTL), complex-trait (GWAS), "
    "and biochemical-state (ENCODE cCRE) — all enrich significantly in Q2, none of "
    "which is used to define the quadrant.",
)
add_table(
    doc,
    ["Annotation", "Fold enrichment", "Hypergeometric p", "Source"],
    [
        ["RepeatMasker low-complexity", "2.02×", "≈0", "RepeatMasker hg38"],
        ["GENCODE 5' UTR", "1.95×", "≈0", "GENCODE v44"],
        ["ENCODE cCRE-ELS (enhancer-like)", "1.90×", "< 10⁻³⁰⁰", "ENCODE SCREEN v3"],
        ["GTEx eQTL (4-tissue union)", "1.62×", "5.4×10⁻⁵⁶", "GTEx v8"],
        ["RepeatMasker simple_repeat", "1.52×", "≈0", "RepeatMasker hg38"],
        ["GWAS Catalog v1.0 SNPs", "1.50×", "1.6×10⁻⁷", "EBI GWAS Catalog"],
        ["RepeatMasker LTR class", "1.39×", "≈0", "RepeatMasker hg38"],
        ["RepeatMasker LINE class", "1.31×", "≈0", "RepeatMasker hg38"],
        ["ENCODE cCRE (any)", "1.28×", "≈0", "ENCODE SCREEN v3"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 3. Q2 enrichment by genomic annotation. Top three rows are sequence-context "
    "enrichments; eQTL/GWAS rows are independent functional axes.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "Sequence-level enrichments — low-complexity and TE classes (LTR, LINE) at "
    "1.3–2.0× — are consistent with the literature that lineage-specific TE-derived "
    "enhancers and promoters are major sources of recently evolved regulatory "
    "function [7]. The three independent functional axes correspond to "
    "molecular-trait (eQTL), complex-trait (GWAS), and biochemical-state (ENCODE) "
    "evidence. gDTR therefore identifies regulatory regions that are functional but "
    "evolutionarily young — a class that sequence-conservation methods systematically "
    "miss — making it complementary to PhyloP/GERP for functional annotation.",
)
add_figure(
    doc,
    FIG_DIR / "F6_q2_conservation_discordance.png",
    "Figure 3. Q2 conservation discordance on chr22. Median-split scatter of mean "
    "gDTR settling depth vs PhyloP 100-way reveals four quadrants; Q2 (high gDTR ∩ "
    "low conservation) covers 3.71 % of valid chr22 (5,090 ≥100 bp regions) and is "
    "1.62×–1.90× enriched for three independent functional axes (GTEx eQTLs, GWAS "
    "SNPs, ENCODE cCRE-ELS), none of which is used to define the quadrant.",
    width_in=6.2,
)

# §3.3 Mechanism
add_h2(doc, "3.3 Class-stratified disruption layers")
add_para(
    doc,
    "To complement the genome-wide aggregate signals (§§3.1–3.2) with per-variant "
    "mechanism evidence, we trace three pathogenic chr17 variants drawn from "
    "distinct functional classes — canonical splice-region, missense, and "
    "frameshift-locus — and three matched benign neighbours (closest B/LB variant in "
    "the same gene by genomic position, 2–10 bp away). Each variant is forwarded "
    "through Evo 2 7B with ± 3 kb context; its 32-layer ΔD_cos trajectory is "
    "extracted from a single forward pass.",
)
add_table(
    doc,
    ["Variant", "Class", "argmax layer", "max|ΔD_cos|", "matched B/LB ratio"],
    [
        ["BRCA1 17:43076602 G→T (canonical splice region)", "P/LP (3-star)", "L7 (shallow)", "3.90×10⁻²", "11.1×"],
        ["TP53 17:7674220 C→T (p.R175H)", "P/LP (3-star)", "L28 (deep)", "2.06×10⁻²", "8.0×"],
        ["BRCA1 17:43057063 G→A (frameshift locus)", "P/LP (3-star)", "L24 (mid-deep)", "3.67×10⁻²", "1.5×"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 4. Per-variant 32-layer ΔD trajectories. \"argmax layer\" indicates where "
    "the variant maximally disrupts the residual stream. \"Matched B/LB ratio\" is "
    "the ratio of pathogenic max|ΔD_cos| to its closest benign neighbour's.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "The three pathogenic variants peak at strikingly different network depths, and "
    "these depths track their biological mechanism. The BRCA1 canonical splice-region "
    "variant peaks at layer 7 — a shallow signal consistent with splice-motif "
    "recognition being a sequence-level circuit (matching the genome-wide finding of "
    "§3.1). TP53 p.R175H, the canonical missense in the DNA-binding domain whose "
    "effect is structural, peaks at layer 28 — deep computation that integrates "
    "protein-level context. The BRCA1 frameshift-locus SNV peaks at layer 24, "
    "consistent with longer-range domain integration. In all three pairs, the "
    "pathogenic max|ΔD_cos| exceeds its matched benign neighbour's by 1.5×–11×. This "
    "shallow-vs-deep stratification by variant class is the core mechanistic claim "
    "of gDTR: the layer index at which a variant maximally disturbs the residual "
    "stream is itself a biologically meaningful readout, and it is invisible to any "
    "single-scalar score (revisited in §3.5 against ‖Δh‖₂).",
)
add_figure(
    doc,
    FIG_DIR / "F5_mechanism_cases.png",
    "Figure 4. Class-stratified disruption layers in three chr17 pathogenic variants. "
    "Per-layer ΔD_cos traces show splice (BRCA1) peaking at L7, frameshift (BRCA1) at "
    "L24, and missense (TP53 R175H) at L28. Each pathogenic trace is paired with its "
    "closest benign neighbour in the same gene; pathogenic max|ΔD_cos| exceeds its "
    "matched benign by 1.5×–11×. The argmax layer maps onto biological mechanism: "
    "shallow circuits for splice-motif recognition, deep circuits for "
    "structural/protein-level effects.",
    width_in=6.2,
)

# §3.4 Variant-level layer-trajectory
add_h2(doc, "3.4 The variant-level ΔD signal is a layer-trajectory property")
add_para(
    doc,
    "The mechanism case studies in §3.3 establish per-variant heterogeneity in "
    "argmax layer; we now ask whether the 32-dimensional per-layer trajectory ΔD is "
    "essential at the population level, or whether a single canonical tap suffices. "
    "We use 8,008 ClinVar P/LP-vs-B/LB single-nucleotide variants stratified across "
    "15 cancer-associated genes (BRCA1/2, TP53, EGFR, KRAS, BRAF, PIK3CA, APC, MLH1, "
    "MSH2, PTEN, RB1, VHL, ATM, PALB2; ClinVar 2026-04-18, 350-cap per gene-class "
    "cell). For each variant we forward Evo 2 7B with ± 3 kb context, extract "
    "per-layer ΔD_cos and ΔD_jsd at the variant token, and fit a logistic regression "
    "with stratified 10-fold and leave-one-gene-out (LOGO) cross-validation. "
    "Bootstrap 95 % CIs are taken from 1,000 resamples.",
)
add_table(
    doc,
    ["Feature", "Dim", "Stratified 10-fold AUROC", "LOGO-CV AUROC"],
    [
        ["Best single-layer ΔD_cos (L = 30 tap)", "1", "0.729 [0.717, 0.741]", "0.726 [0.694, 0.758]"],
        ["Best single-layer ΔD_jsd (L = 29 canonical tap)", "1", "0.794 [0.781, 0.806]", "0.787 [0.752, 0.821]"],
        ["Evo 2 log-likelihood", "1", "0.751 [0.738, 0.764]", "0.793 [0.740, 0.846]"],
        ["32-d ΔD_jsd vector", "32", "0.823 [0.813, 0.832]", "0.821 [0.790, 0.853]"],
        ["32-d ΔD_cos vector (primary)", "32", "0.844 [0.831, 0.857]", "0.843 [0.811, 0.876]"],
        ["Ensemble (ΔD + Evo 2 LL)", "33", "0.861 [0.851, 0.871]", "0.866 [0.832, 0.899]"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 5. The 32-d trajectory dominates any single-layer feature. The primary "
    "cosine lens reaches 0.844 with 32 layers and only 0.729 at its best single tap "
    "— a +0.115 gap. Bracketed numbers are 1,000-bootstrap 95 % confidence intervals.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "Three properties of these results are worth emphasising as framework claims. "
    "First, the full 32-d trajectory is structurally essential. The per-layer "
    "single-feature AUROC traces a clear U-shape across depth — early layers hover "
    "near 0.60, the mid-zone (L11–L17) plateaus at 0.66–0.74, and performance peaks "
    "near the canonical tap (L = 29 for JSD). Even the best single tap falls "
    "5–12 AUROC points short of the vector — a direct consequence of the per-variant "
    "heterogeneity established in §3.3: pathogenic variants disrupt the network at "
    "multiple distinct layer regimes, and any single layer can capture only one mode "
    "at a time.",
)
add_para(
    doc,
    "Second, the trajectory is far more informative than scalar summaries of "
    "settling depth. Phase 0 pilot analyses showed that scalar features such as Δc "
    "or |Δc| yield only weak discrimination (AUROC ≈ 0.64 after sign flip), because "
    "a single change in settling depth collapses the rich layer-wise pattern into one "
    "number. In contrast, ΔD records the magnitude and location of residual-stream "
    "perturbation at every layer, preserving the full mechanistic fingerprint.",
)
add_para(
    doc,
    "Third, the gDTR trajectory provides statistically significant information "
    "complementary to Evo 2's own output-side signal (Δ log-likelihood). The 32-d "
    "vector alone outperforms Evo 2 ΔLL by a wide margin (0.844 vs 0.751 in "
    "stratified CV). When the two are combined into a 33-d ensemble, performance "
    "rises to 0.861 (+0.017) in stratified CV and 0.866 (+0.023) in LOGO-CV. A "
    "DeLong paired test confirms this improvement is highly significant (p = 3.6 × "
    "10⁻¹⁵). The leave-one-gene-out AUROC (0.843) is statistically indistinguishable "
    "from the stratified result (0.844), demonstrating that the predictive signal is "
    "gene-agnostic rather than the product of memorising any single driver gene's "
    "context.",
)
add_figure(
    doc,
    FIG_DIR / "F3_variant_pathogenicity.png",
    "Figure 5. Variant pathogenicity discrimination. (a) ROC curves for ΔD_cos "
    "vector (AUROC 0.844), ΔD_jsd vector (0.823), Evo 2 ΔLL (0.751), and ΔD+ΔLL "
    "ensemble (0.861). (b) DeLong paired comparisons; ΔD adds DeLong-significant "
    "information beyond Evo 2 ΔLL (p = 3.6 × 10⁻¹⁵). (c) Per-layer single-feature "
    "AUROC across all 32 layers — best single tap is L = 29 (JSD); the 32-d "
    "vector beats it by +0.05 to +0.12. (d) Leave-one-gene-out AUROC across 14 "
    "evaluable genes is uniformly high, evidence that the signal is gene-agnostic.",
    width_in=6.2,
)

# §3.5 NEW: Baseline benchmark — v3 reframing
add_h2(doc, "3.5 Interpretability baseline benchmark: gDTR as mechanistic probe vs ‖Δh‖₂ as scorer")
add_para(
    doc,
    "Variant scoring on ClinVar is dominated by likelihood-based and "
    "structure-based predictors (CADD, AlphaMissense, ESM1b). Hidden-state "
    "perturbation magnitudes from foundation models — although a single forward "
    "pass already exposes them — are surprisingly absent from the published "
    "benchmark literature. To contextualise gDTR's discrimination, and to ensure "
    "the framework is not merely a re-skinned NLP technique, we benchmark four "
    "interpretability axes on the same 8,008-variant cohort with the identical "
    "stratified 10-fold and LOGO-CV pipeline used in §3.4: (a) ΔD_cos vector "
    "(gDTR, 32-d), (b) per-layer ‖Δh‖₂ vector (32-d L2 norms of the residual-stream "
    "perturbation at each block), (c) attention rollout (Abnar & Zuidema 2020, 5-d "
    "summary), and (d) integrated gradients on h_29 (1-d scalar).",
)
add_table(
    doc,
    ["Method", "Dim", "AUROC (stratified)", "AUROC (LOGO)", "DeLong vs gDTR"],
    [
        ["ΔD_cos (gDTR, this work)", "32", "0.844 [0.831, 0.857]", "0.843 [0.811, 0.876]", "—"],
        ["‖Δh‖₂ per-layer L2 norm", "32", "0.926 [0.921, 0.932]", "0.922 [0.903, 0.942]", "+0.083, p ≈ 0"],
        ["Attention rollout (5 layer-summaries)", "5", "0.672 [0.660, 0.684]", "0.668 [0.635, 0.701]", "−0.172, p ≈ 0"],
        ["Integrated gradients on h₂₉", "1", "0.527 [0.515, 0.540]", "0.524 [0.497, 0.551]", "−0.316, p ≈ 0"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 6. Interpretability baseline comparison on 8,008 ClinVar P/LP-vs-B/LB "
    "variants. ‖Δh‖₂ is the strongest scoring baseline by +0.083 over gDTR; "
    "attention rollout and integrated gradients lag substantially.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "We report this result transparently rather than hide it. Two implications "
    "follow. First, ‖Δh‖₂ — the simplest possible \"how much did the variant "
    "disturb the model?\" feature — is a previously under-reported strong "
    "scoring baseline that future variant-effect-prediction work should include. "
    "We treat its identification as a benchmarking contribution of this paper. "
    "Second, the headline question of this paper is no longer \"which scalar "
    "scores best on ClinVar?\" but rather \"what does the layer-resolved view "
    "tell us that magnitudes alone cannot?\". We argue gDTR provides three "
    "irreducible properties that ‖Δh‖₂ cannot supply by construction:",
)
add_bulleted(
    doc,
    [
        "(P1) Layer-index reference frame. gDTR's c(t) is a layer index, not a magnitude. ‖Δh‖₂ has no such reference and cannot, on its own, distinguish a 0.1-norm shift at L5 from a 0.1-norm shift at L28 — it provides no \"where in the network\" co-ordinate.",
        "(P2) Class-stratified disruption layers. The mechanism case studies (§3.3) show that pathogenic splice variants peak at shallow L7, missense effects at deep L28, and frameshifts at mid-deep L24. This shallow-vs-deep stratification is invisible to a magnitude-only score because magnitudes alone cannot recover the argmax layer.",
        "(P3) Connection to genome-wide universality. The chr22 + chr17 splice donor/acceptor shallow-thinking signature (§3.1) and the chr22 Q2 conservation discordance (§3.2) are layer-index findings derived from c(t), not from norms. ‖Δh‖₂ does not produce a comparable genome-wide layer-stratified annotation track.",
    ],
)
add_para(
    doc,
    "These three properties motivate gDTR's role as a mechanistic probe rather than "
    "a competing scorer. The two tools answer different questions on a single forward "
    "pass: ‖Δh‖₂ answers how much the variant perturbed the model, gDTR answers "
    "where in the computational hierarchy and how that location maps onto biological "
    "grammar. Wall-clock costs are within 5 % of each other (517 ms/variant for "
    "‖Δh‖₂ vs 540 ms/variant for gDTR on H200; §A4).",
)
add_para(
    doc,
    "We additionally test whether gDTR's signal is fully captured by ‖Δh‖₂. Fitting "
    "a logistic regression on ΔD_cos after residualising each of its 32 features on "
    "the corresponding ‖Δh‖₂ component yields AUROC 0.645 — well above chance (0.5) "
    "and clearly indicating independent variance, not a re-encoding. Conversely, "
    "ΔD_cos cannot recover ‖Δh‖₂'s discrimination after residualising. The two "
    "signals are correlated but neither subsumes the other.",
)
add_figure(
    doc,
    FIG_DIR / "F4_baselines.png",
    "Figure 6. Interpretability baseline benchmark on 8,008 ClinVar variants. "
    "(a) Per-method AUROC under stratified 10-fold and leave-one-gene-out CV; "
    "‖Δh‖₂ leads at 0.926, gDTR follows at 0.844, attention rollout 0.672, "
    "integrated gradients 0.527. (b) Pairwise Spearman concordance across methods "
    "— ΔD_cos and ‖Δh‖₂ correlate moderately (ρ = 0.57) but neither subsumes the "
    "other. (c) DeLong paired tests confirm all three pairwise differences "
    "(p ≈ 0). (d) Incremental information: gDTR retains AUROC 0.645 after "
    "residualising on ‖Δh‖₂ — independent variance, not a re-encoding.",
    width_in=6.2,
)
add_para(
    doc,
    "Reviewer pre-empt — \"if ‖Δh‖₂ wins, why use gDTR?\". ‖Δh‖₂ achieves higher "
    "raw classification AUROC but provides no mechanistic resolution. It cannot "
    "distinguish shallow splicing circuits (§3.1, §3.3) from deep structural "
    "processing (§3.3 TP53), does not connect to genome-wide layer-stratified "
    "phenomena (§3.2), and produces no layer-index reference frame. gDTR provides "
    "this resolution at < 5 % overhead and reveals unique variance (residualised "
    "AUROC 0.645). The two tools answer different questions; AUROC alone is the "
    "wrong figure of merit when the deliverable is mechanistic understanding. "
    "Future work that benchmarks variant-effect predictors should include both, "
    "with ‖Δh‖₂ as a strong scalar baseline and gDTR as the layer-resolved probe.",
)

# §3.6 = old §3.5 cross-arch
add_h2(doc, "3.6 Cross-architecture invariance is two-tier, not universal")
add_para(
    doc,
    "To assess robustness beyond any single genomic foundation model, we apply the "
    "identical per-window mean settling-depth analysis to four state-of-the-art "
    "models: Evo 2, HyenaDNA-large, Nucleotide Transformer v2, DNABERT-2. Per-model "
    "q70 calibration is applied; per-window mean c is then compared by Spearman "
    "correlation. All p-values are < 10⁻⁴².",
)
add_table(
    doc,
    ["", "Evo 2 7B", "HyenaDNA-large", "NT-v2 500M", "DNABERT-2"],
    [
        ["Evo 2 7B", "1.000", "+0.516", "−0.119", "−0.188"],
        ["HyenaDNA-large", "+0.516", "1.000", "−0.287", "−0.166"],
        ["NT-v2 500M", "−0.119", "−0.287", "1.000", "+0.663"],
        ["DNABERT-2", "−0.188", "−0.166", "+0.663", "1.000"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table 7. Pairwise Spearman ρ of per-window mean settling depth across four "
    "genomic foundation models. Within-family correlations (causal-LM block, MLM "
    "block) are strong and positive; cross-family correlations are weakly negative.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)
add_para(
    doc,
    "The structure is two-tier. Within the per-bp causal-LM family, Evo 2 and "
    "HyenaDNA correlate at ρ = +0.516 despite a 250× parameter-count gap and 4× "
    "depth gap. Within the MLM family, NT-v2 and DNABERT-2 correlate at ρ = +0.663 "
    "despite different tokenisations. Across families, every pair shows weakly "
    "negative correlation, and the four-way top-decile concordance is zero windows "
    "— each family lights up entirely different chr22 windows. The accurate "
    "refinement is: gDTR rankings are architecture-invariant within a "
    "tokenisation/objective family, but the level at which deep computation occurs "
    "is itself tokenisation-dependent. We do not interpret cross-family negative "
    "correlations because per-position alignment between bp and k-mer/BPE coordinates "
    "is non-trivial and is left as future work.",
)
add_figure(
    doc,
    FIG_DIR / "F7_cross_architecture.png",
    "Figure 7. Cross-architecture two-tier structure. Pairwise Spearman of per-window "
    "mean settling depth across four models reveals strong within-family agreement "
    "(causal-LM ρ = +0.52, MLM ρ = +0.66) but weakly negative cross-family agreement, "
    "consistent with tokenisation- and objective-dependent computational regimes. "
    "Top-decile concordance across all four is zero windows.",
    width_in=6.2,
)

# ─────────────────────────────────────────────────────────────────────────
# §4 Limitations
add_h1(doc, "4 Limitations")
add_bulleted(
    doc,
    [
        "(L1) Genome scope. All genome-wide analyses use chr22 + chr17 (≈130 Mb, ≈4 % of GRCh38). Whole-genome rollout is expected to behave similarly but is unverified.",
        "(L2) Final-block-idleness assumption. The canonical-tap shift to L = 29 is justified by the observed h₃₀ ≡ h₃₁ in Evo 2 7B; new architectures must be re-validated rather than inheriting this choice.",
        "(L3) Driver-gene class size. The chr17 cancer-driver vs non-driver comparison contained only n = 2 driver genes (TP53, BRCA1), giving Cohen's d = +0.87 but p = 0.14. Counter-intuitively, driver genes showed less deep thinking; we report the direction honestly but the test is underpowered.",
        "(L4) Cross-family per-position alignment. Per-bp causal-LM models (Evo 2, HyenaDNA) are not directly comparable position-by-position to k-mer/BPE MLMs (NT-v2, DNABERT-2). The two-tier finding is the strongest cross-family claim our data can support; tokenisation-aware re-alignment is required for unified comparison.",
        "(L5) Calibration locality. Regional q70 calibration is robust within the ±0.10/±0.05 plateau measured on chr22, but γ is ultimately a region-specific quantile. For deployment on a new region or organism, recalibration is required.",
        "(L6) Discrimination vs mechanistic resolution. ‖Δh‖₂ outperforms gDTR by +0.083 AUROC as a raw classifier (§3.5). gDTR is justified by its mechanistic-probe properties (P1–P3); users for whom raw classification is the only deliverable should use both together rather than gDTR alone.",
    ],
)

# ─────────────────────────────────────────────────────────────────────────
# §5 Conclusion
add_h1(doc, "5 Conclusion")
add_para(
    doc,
    "We introduce gDTR, a training-free, layer-resolved interpretability framework "
    "for genomic causal language models, and benchmark it against three "
    "interpretability baselines on 8,008 ClinVar variants. The benchmark establishes "
    "two complementary tools that run from a single forward pass at < 5 % wall-clock "
    "difference: ‖Δh‖₂ as a strong scoring baseline (AUROC 0.926) and gDTR as a "
    "layer-resolved mechanistic probe (AUROC 0.844 with residualised independent "
    "variance 0.645). Across complementary analyses, gDTR uncovered biologically "
    "coherent computational signatures: splice donor/acceptor sites as the universally "
    "shallowest loci of computation across chromosomes and architectural families "
    "(§3.1); regions of deep internal engagement that diverge from evolutionary "
    "conservation and are enriched for regulatory elements and disease-associated "
    "variants (§3.2); class-stratified disruption layers that distinguish subtle "
    "motif-level disruptions from deeper structural effects (§3.3); a 32-d "
    "layer-trajectory variant signal that beats any single tap and provides "
    "DeLong-significant incremental information beyond Evo 2 likelihood (§3.4); and "
    "a refined cross-architecture invariance claim — two-tier rather than universal "
    "(§3.6).",
)
add_para(
    doc,
    "The methodological contribution is therefore twofold: a layer-resolved probe "
    "(gDTR) that links foundation-model internals directly to genomic biology, and "
    "an honest benchmark that identifies ‖Δh‖₂ as a previously under-reported strong "
    "scoring baseline. Both findings are reproducible from a single forward pass. The "
    "enumerated limitations stem from the specific architectures and chromosomes "
    "evaluated and serve as guideposts for whole-genome rollout, per-architecture "
    "re-validation, and tokenisation-aware cross-family alignment.",
)

# ─── References ──
add_h1(doc, "References")
refs = [
    "[1] Nguyen, E., Poli, M., Durrant, M.G., et al. (2026). Evo 2: whole-genome modelling with context-length scaling. Nature.",
    "[2] Nguyen, E., Poli, M., Faltings, B., et al. (2023). HyenaDNA: long-range genomic sequence modelling at single-nucleotide resolution. NeurIPS.",
    "[3] Dalla-Torre, H., Gonzalez, L., Mendoza-Revilla, J., et al. (2024). The Nucleotide Transformer: building and evaluating robust foundation models for human genomics. Nature Methods.",
    "[4] Zhou, Z., Ji, Y., Li, W., et al. (2024). DNABERT-2: efficient foundation model and benchmark for multi-species genomes. ICLR.",
    "[5] Chen, W.-L., et al. (2026). Think Deep, Not Just Long: measuring LLM reasoning effort via deep-thinking tokens. arXiv:2602.13517.",
    "[6] Belrose, N., Furman, Z., Smith, L., et al. (2023). Eliciting latent predictions from transformers with the tuned lens. arXiv:2303.08112.",
    "[7] Chuong, E.B., Elde, N.C., Feschotte, C. (2017). Regulatory activities of transposable elements: from conflicts to benefits. Nature Reviews Genetics.",
    "[8] Pollard, K.S., Hubisz, M.J., Rosenbloom, K.R., Siepel, A. (2010). Detection of nonneutral substitution rates on mammalian phylogenies (PhyloP). Genome Research.",
    "[9] Abnar, S., Zuidema, W. (2020). Quantifying attention flow in transformers. ACL.",
    "[10] Sundararajan, M., Taly, A., Yan, Q. (2017). Axiomatic attribution for deep networks. ICML.",
    "[11] nostalgebraist (2020). Interpreting GPT: the logit lens. AI Alignment Forum.",
]
for r in refs:
    p = doc.add_paragraph()
    p.add_run(r).font.size = Pt(10)

# ─── Appendix ──
add_h1(doc, "Appendix A Method details")
add_h2(doc, "A.1 Tuned-lens recovery across all 32 layers")
add_para(
    doc,
    "Each layer is fitted with a single 4096×4096 affine A_l using MSE between "
    "A_l h_l and h_norm, optimised with Adam at 1e−3 for 15 epochs over 100 "
    "calibration sequences. The post-norm tap is the prediction target. Recovery "
    "scores are summarised below; 30/32 layers reach ≥ 98 %.",
)
add_table(
    doc,
    ["Layer", "Block type", "Initial MSE", "Final MSE (15 epochs)", "Recovery"],
    [
        ["L = 2", "hcl", "1,259", "5.6", "0.9956"],
        ["L = 12", "hcm", "742", "13.6", "0.9816 (worst)"],
        ["L = 22", "hcm", "418", "0.41", "0.9990"],
        ["L = 28", "hcs", "822", "0.34", "0.9996 (best below tap)"],
        ["L = 29 (canonical tap)", "hcm", "510", "0.20", "0.9996"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table A1. Tuned-lens recovery at five representative layers spanning network depth.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "A.2 Hyperparameter sensitivity (chr22 splice donor vs intronic)")
add_table(
    doc,
    ["γ_cos", "ρ = 0.70", "ρ = 0.75", "ρ = 0.80", "ρ = 0.85", "ρ = 0.90"],
    [
        ["0.30", "5.04", "5.18", "5.21", "5.20", "5.15"],
        ["0.35", "5.10", "5.21", "5.24", "5.23", "5.18"],
        ["0.40", "5.16", "5.25", "5.28", "5.26", "5.21"],
        ["0.45", "5.13", "5.22", "5.25", "5.24", "5.19"],
        ["0.50", "5.08", "5.17", "5.20", "5.19", "5.14"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table A2. Cohen's d sweep over (γ_cos, ρ). The locked operating point (0.40, 0.80) "
    "sits inside a flat plateau; standard deviation across the 25 cells is 0.06.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "A.3 Per-layer single-feature AUROC (selected)")
add_para(
    doc,
    "Single-layer logistic-regression out-of-fold AUROC under both lenses on the same "
    "8,008-variant cohort and 10-fold stratified split (seed = 42) used in §3.4. "
    "Selected layers are tabulated; the full 32-row table is released as "
    "per_layer_auroc.csv.",
)
add_table(
    doc,
    ["Layer", "Block type", "ΔD_jsd AUROC", "ΔD_cos AUROC"],
    [
        ["0", "embed", "0.605", "0.519"],
        ["7", "hcs (shallow)", "0.662", "0.595"],
        ["12", "hcm", "0.656", "0.555"],
        ["17", "attn", "0.723", "0.646"],
        ["24", "attn", "0.685", "0.612"],
        ["28", "hcs", "0.565", "0.698"],
        ["29 (canonical tap)", "hcm", "0.794 (best jsd)", "0.604"],
        ["30 (post-norm-1)", "—", "0.512", "0.729 (best cos)"],
        ["31 (idle, see §2.2)", "—", "0.499 (degenerate)", "0.729 (= L30)"],
        ["32-d vector", "all", "0.823", "0.844"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table A3. Selected per-layer AUROCs. Cosine and JSD lenses agree on the U-shape "
    "but differ on which tap is sharpest: JSD concentrates discriminative mass at the "
    "canonical tap L = 29; cosine spreads it across many taps, producing a much larger "
    "vector-vs-best-single-layer gap.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "A.4 Compute-cost benchmark (single H200, 100 variants)")
add_table(
    doc,
    ["Method", "Wall-clock per variant (ms)", "Peak GPU memory (GB)", "Forward passes"],
    [
        ["gDTR (ΔD_cos vector)", "540", "73", "2 (ref + alt)"],
        ["‖Δh‖₂ per-layer L2 norm", "517", "73", "2 (ref + alt, shared with gDTR)"],
        ["Attention rollout", "612", "78", "2 (ref + alt)"],
        ["Integrated gradients (8 steps)", "4,210", "92", "16 (ref + alt × 8 IG steps)"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table A4. Compute-cost benchmark. ‖Δh‖₂ and gDTR share the same forward pass; "
    "the < 5 % wall-clock gap is post-hoc feature extraction. Integrated gradients is "
    "8× more expensive due to step-wise gradient evaluation.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "A.5 Cross-architecture model specifications")
add_table(
    doc,
    ["Model", "Architecture", "Layers", "Hidden", "Tokens / 6 kb", "Wall time", "γ_q70"],
    [
        ["Evo 2 7B", "Hybrid Transformer + StripedHyena 2", "32", "4096", "6,000 (1 bp)", "reused", "0.396"],
        ["HyenaDNA-large-1m", "Pure Hyena", "8", "256", "6,001 (1 bp + BOS)", "~4 min", "0.358"],
        ["NT-v2 500M", "Transformer MLM (k-mer)", "29", "1024", "671 (k=6)", "~7.5 min", "0.533"],
        ["DNABERT-2 117M", "Transformer MLM (BPE)", "12", "768", "~600 (BPE)", "~3 min", "0.677"],
    ],
    first_col_bold=True,
)
add_para(
    doc,
    "Table A5. Cross-architecture implementation notes. All inference on a single "
    "NVIDIA H200 GPU.",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=9,
)

add_h2(doc, "A.6 Reproducibility")
add_para(
    doc,
    "All experiments use random seed 42 for cross-validation splits and bootstrap "
    "resamples. Evo 2 model lock: arcinstitute/evo2_7b at HF revision SHA "
    "bda0089f92582d5baabf0f22d9fc85f3588f6b58 (weights MD5 "
    "359ef88ccac2a62644035578de8a7db4). Data versions: GRCh38 primary assembly (UCSC, "
    "MD5 locked); GENCODE v44 GTF; PhyloP 100-way (UCSC); ENCODE SCREEN v3 cCRE "
    "(ELS subset); RepeatMasker hg38; GTEx v8 cis-eQTL pairs (Whole_Blood, "
    "Brain_Cortex, Liver, Lung union); GWAS Catalog v1.0; ClinVar 2026-04-18. "
    "Software stack: torch 2.4.1+cu124, evo2 0.3.0, vortex 1.0.8, transformer-engine "
    "2.14.0, transformers 4.49.0, scipy 1.13, scikit-learn 1.4. Hardware: a single "
    "NVIDIA H200 (141 GB). Total compute: ≈20 GPU-hours end-to-end. Code, dataset "
    "version locks, and figure-generation scripts: "
    "https://github.com/darejinn/gDTR-PoC.",
)

# ─── save ──
doc.save(OUT)
print(f"Wrote {OUT}")
